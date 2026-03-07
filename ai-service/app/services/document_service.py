import io
import structlog

logger = structlog.get_logger()


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    logger.info("docx_extracted", chars=len(text))
    return text


def extract_text_from_xlsx(xlsx_bytes: bytes) -> str:
    """Extract text from XLSX file, formatting as a readable table."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
    lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"=== Лист: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            values = [str(v) if v is not None else "" for v in row]
            if any(v.strip() for v in values):
                lines.append(" | ".join(values))

    wb.close()
    text = "\n".join(lines)
    logger.info("xlsx_extracted", chars=len(text))
    return text
