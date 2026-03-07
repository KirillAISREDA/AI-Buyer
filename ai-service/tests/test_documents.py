"""Tests for document extraction services."""
import io
from app.services.document_service import extract_text_from_xlsx, extract_text_from_docx


def test_extract_xlsx():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Наименование", "Количество", "Цена"])
    ws.append(["Бумага А4", 10, 350])
    ws.append(["Ручка", 50, 25])

    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    text = extract_text_from_xlsx(xlsx_bytes)
    assert "Бумага А4" in text
    assert "Ручка" in text
    assert "350" in text


def test_extract_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Счёт №123 от 15.01.2026")
    doc.add_paragraph("Поставщик: ООО Ромашка")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Товар"
    table.rows[0].cells[1].text = "Кол-во"
    table.rows[0].cells[2].text = "Цена"
    table.rows[1].cells[0].text = "Бумага А4"
    table.rows[1].cells[1].text = "10"
    table.rows[1].cells[2].text = "350"

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    text = extract_text_from_docx(docx_bytes)
    assert "Счёт №123" in text
    assert "ООО Ромашка" in text
    assert "Бумага А4" in text


def test_extract_xlsx_empty_rows():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append([None, None, None])
    ws.append(["Данные", 1, 100])

    buf = io.BytesIO()
    wb.save(buf)

    text = extract_text_from_xlsx(buf.getvalue())
    assert "Данные" in text
